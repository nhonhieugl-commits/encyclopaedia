"""
Bách Khoa Toàn Thư Ops — AI Agent
Internal Knowledge Base for Merchant Operations Team

Stack: FastAPI + OpenAI-compatible SDK (GreenNode MaaS) + SQLite + JWT Auth
Deploy: GreenNode AgentBase (HTTP runtime, port 3000)
"""

import os
import json
import sqlite3
import hashlib
import secrets
import logging
import asyncio
from datetime import datetime, timedelta
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Optional, AsyncGenerator

from dotenv import load_dotenv
load_dotenv()  # Load .env file nếu có

from openai import OpenAI
import anthropic as anthropic_sdk
import jwt
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form, Request, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, HTMLResponse
from pydantic import BaseModel

FRONTEND_PATH = Path(__file__).parent / "frontend.html"

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("ops-agent")

API_KEY = os.getenv("ANTHROPIC_API_KEY", "")  # Dùng chung env var, tương thích GreenNode key
API_BASE_URL = os.getenv("ANTHROPIC_BASE_URL", "http://maas.greennode.ai/v1")  # GreenNode MaaS endpoint
JWT_SECRET = os.getenv("JWT_SECRET", "CHANGE_THIS_TO_RANDOM_SECRET_MIN_32_CHARS")
JWT_EXPIRES_IN = os.getenv("JWT_EXPIRES_IN", "8h")
MODEL = os.getenv("MODEL", "Qwen/Qwen3-30B-A3B-FP8")
MAX_TOKENS = int(os.getenv("MAX_TOKENS", "4096"))
PORT = int(os.getenv("PORT", "3000"))
DB_PATH = os.getenv("DB_PATH", "/app/data/ops.db")
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "/app/data/uploads"))
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "20"))
VERSION = "1.0.0"

SECTION_NAMES = {
    1: "Từ Điển Thuật Ngữ Ops",
    2: "Quy Trình Chuẩn",
    3: "Cẩm Nang Xử Lý Sự Cố",
    4: "Danh Mục Thiết Bị & Phần Cứng",
    5: "Kiểm Tra Kiến Thức",
}

SYSTEM_PROMPT = """Bạn là trợ lý AI nội bộ của đội Merchant Operations, được xây dựng trên nền tảng kiến thức của "Bách Khoa Toàn Thư Ops".

Nhiệm vụ của bạn:
- Trả lời câu hỏi về quy trình, thuật ngữ, thiết bị và xử lý sự cố dựa trên knowledge base
- Hướng dẫn nhân viên thực hiện đúng quy trình chuẩn (SOP)
- Giải thích các thuật ngữ kỹ thuật và nghiệp vụ một cách rõ ràng
- Hỗ trợ xử lý sự cố thường gặp với Soundbox, cổng thanh toán, kết nối thiết bị

Nguyên tắc:
- Luôn dùng tiếng Việt, trừ khi được yêu cầu khác
- Chỉ trả lời dựa trên thông tin có trong knowledge base (dùng tools để tra cứu)
- Nếu không tìm thấy thông tin, nói rõ và gợi ý liên hệ admin để cập nhật tài liệu
- Trả lời ngắn gọn, có cấu trúc rõ ràng, dùng bullet points khi liệt kê các bước
- Không đưa ra thông tin tài chính, pháp lý, hoặc thông tin bảo mật

5 phần nội dung chính:
1. Từ Điển Thuật Ngữ Ops — TPV, Lead, CR%, Onboard, Churn...
2. Quy Trình Chuẩn — SOPs, Memo, chính sách nội bộ
3. Cẩm Nang Xử Lý Sự Cố — Soundbox, lỗi kết nối, cổng thanh toán
4. Danh Mục Thiết Bị & Phần Cứng — Hướng dẫn sử dụng & reset
5. Kiểm Tra Kiến Thức — Quiz trắc nghiệm"""

# Tool definitions — OpenAI function-calling format (tương thích GreenNode MaaS)
TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "search_knowledge_base",
            "description": "Tìm kiếm tài liệu trong knowledge base theo từ khóa. Dùng khi người dùng hỏi về thuật ngữ, quy trình, hoặc thiết bị cụ thể.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Từ khóa tìm kiếm (tiếng Việt hoặc tiếng Anh)"
                    },
                    "section": {
                        "type": "integer",
                        "description": "Lọc theo section 1-5. Bỏ qua để tìm toàn bộ.",
                        "minimum": 1,
                        "maximum": 5
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Số kết quả tối đa (mặc định 5)",
                        "default": 5
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_section_content",
            "description": "Lấy danh sách tài liệu trong một section cụ thể. Dùng khi người dùng muốn xem tổng quan một phần.",
            "parameters": {
                "type": "object",
                "properties": {
                    "section": {
                        "type": "integer",
                        "description": "Section ID (1-5)",
                        "minimum": 1,
                        "maximum": 5
                    }
                },
                "required": ["section"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_quiz_questions",
            "description": "Lấy câu hỏi quiz để kiểm tra kiến thức. Dùng khi người dùng muốn ôn luyện hoặc làm bài kiểm tra.",
            "parameters": {
                "type": "object",
                "properties": {
                    "section": {
                        "type": "integer",
                        "description": "Section ID (1-5). Bỏ qua để lấy câu hỏi từ tất cả sections.",
                        "minimum": 1,
                        "maximum": 5
                    },
                    "count": {
                        "type": "integer",
                        "description": "Số câu hỏi (mặc định 10, tối đa 20)",
                        "default": 10,
                        "maximum": 20
                    }
                },
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_document_detail",
            "description": "Lấy nội dung đầy đủ của một tài liệu theo ID. Dùng sau search để đọc chi tiết.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {
                        "type": "integer",
                        "description": "ID của tài liệu"
                    }
                },
                "required": ["document_id"]
            }
        }
    }
]

# ---------------------------------------------------------------------------
# Database
# ---------------------------------------------------------------------------

def get_db_path() -> str:
    """Return DB path, creating parent dirs if needed."""
    path = Path(DB_PATH)
    path.parent.mkdir(parents=True, exist_ok=True)
    return str(path)


def get_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(get_db_path(), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


def init_db():
    """Initialize DB schema and seed default users."""
    conn = get_connection()
    c = conn.cursor()

    c.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT DEFAULT 'user' CHECK(role IN ('admin', 'user')),
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            section INTEGER NOT NULL CHECK(section BETWEEN 1 AND 5),
            title TEXT NOT NULL,
            content TEXT NOT NULL,
            file_type TEXT DEFAULT 'md',
            filename TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
            title, content,
            content=documents,
            content_rowid=id,
            tokenize='unicode61'
        );

        CREATE TABLE IF NOT EXISTS quiz_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            section INTEGER NOT NULL CHECK(section BETWEEN 1 AND 5),
            question TEXT NOT NULL,
            option_a TEXT NOT NULL,
            option_b TEXT NOT NULL,
            option_c TEXT NOT NULL,
            option_d TEXT NOT NULL,
            correct_answer TEXT NOT NULL CHECK(correct_answer IN ('A','B','C','D')),
            explanation TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS quiz_attempts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER REFERENCES users(id),
            section INTEGER,
            score INTEGER NOT NULL,
            total INTEGER NOT NULL,
            detail TEXT,
            taken_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS chat_sessions (
            id TEXT PRIMARY KEY,
            user_id INTEGER REFERENCES users(id),
            messages TEXT DEFAULT '[]',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        );

        -- Trigger to sync FTS on insert
        CREATE TRIGGER IF NOT EXISTS documents_ai AFTER INSERT ON documents BEGIN
            INSERT INTO documents_fts(rowid, title, content)
            VALUES (new.id, new.title, new.content);
        END;

        -- Trigger to sync FTS on update
        CREATE TRIGGER IF NOT EXISTS documents_au AFTER UPDATE ON documents BEGIN
            INSERT INTO documents_fts(documents_fts, rowid, title, content)
            VALUES ('delete', old.id, old.title, old.content);
            INSERT INTO documents_fts(rowid, title, content)
            VALUES (new.id, new.title, new.content);
        END;

        -- Trigger to sync FTS on delete
        CREATE TRIGGER IF NOT EXISTS documents_ad AFTER DELETE ON documents BEGIN
            INSERT INTO documents_fts(documents_fts, rowid, title, content)
            VALUES ('delete', old.id, old.title, old.content);
        END;
    """)

    # Seed default users if not exist
    admin_pw = _hash_password("Admin@123")
    user_pw = _hash_password("User@123")

    c.execute(
        "INSERT OR IGNORE INTO users (username, password_hash, role) VALUES (?, ?, ?)",
        ("admin", admin_pw, "admin")
    )
    c.execute(
        "INSERT OR IGNORE INTO users (username, password_hash, role) VALUES (?, ?, ?)",
        ("user", user_pw, "user")
    )

    # Seed sample documents if DB is fresh
    count = c.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
    if count == 0:
        _seed_sample_data(c)

    conn.commit()
    conn.close()
    log.info("Database initialized at %s", get_db_path())


def _seed_sample_data(c: sqlite3.Cursor):
    """Insert sample knowledge base entries."""
    docs = [
        (1, "TPV (Total Payment Volume)", """## TPV — Tổng Doanh Số Thanh Toán

**Định nghĩa:** TPV (Total Payment Volume) là tổng giá trị giao dịch thanh toán được xử lý qua hệ thống trong một khoảng thời gian nhất định.

**Công thức:**
```
TPV = Tổng số giao dịch × Giá trị trung bình mỗi giao dịch
```

**Ứng dụng:**
- Đo lường hiệu quả của merchant
- KPI chính cho đội Merchant Operations
- Cơ sở tính phí dịch vụ

**Ví dụ:** Merchant A xử lý 500 giao dịch/ngày, mỗi giao dịch trung bình 200.000đ → TPV ngày = 100 triệu đồng
"""),
        (1, "CR% (Conversion Rate)", """## CR% — Tỷ Lệ Chuyển Đổi

**Định nghĩa:** Conversion Rate là tỷ lệ phần trăm khách hàng tiềm năng (Lead) chuyển đổi thành merchant thực sự kích hoạt và giao dịch.

**Công thức:**
```
CR% = (Số merchant đã kích hoạt / Tổng số Lead) × 100%
```

**Benchmark ngành:** CR% trung bình 15–25% là tốt.

**Các yếu tố ảnh hưởng:**
- Chất lượng Lead đầu vào
- Thời gian onboarding
- Chất lượng hỗ trợ kỹ thuật
- Chính sách ưu đãi
"""),
        (2, "Quy Trình Onboarding Merchant Mới", """## SOP: Onboarding Merchant Mới

**Thời gian hoàn thành:** Tối đa 3 ngày làm việc

### Bước 1 — Tiếp nhận hồ sơ (Ngày 1)
- Nhận đầy đủ giấy tờ: ĐKKD, CCCD chủ sở hữu, hợp đồng
- Kiểm tra tính hợp lệ trên hệ thống CRM
- Tạo merchant profile, gán mã merchant ID

### Bước 2 — Cài đặt thiết bị (Ngày 1–2)
- Giao thiết bị Soundbox hoặc mPOS
- Cài đặt app, kết nối merchant ID
- Test giao dịch thử: chạy 1 giao dịch 1.000đ
- Xác nhận nhận giao dịch thành công

### Bước 3 — Training (Ngày 2)
- Hướng dẫn sử dụng thiết bị
- Giải thích báo cáo doanh số
- Cung cấp hotline hỗ trợ: 1900 xxxx

### Bước 4 — Hoàn tất (Ngày 3)
- Cập nhật trạng thái "Active" trên CRM
- Gửi email xác nhận + tài liệu hướng dẫn
- Schedule follow-up sau 7 ngày

**Escalation:** Nếu quá 3 ngày chưa hoàn thành, báo cáo Team Lead ngay.
"""),
        (3, "Xử Lý Lỗi Soundbox Không Kết Nối", """## Cẩm Nang: Soundbox Không Kết Nối / Không Phát Âm Thanh

### Triệu chứng
- Soundbox không phát âm thanh xác nhận giao dịch
- App báo "Thiết bị offline"
- Đèn LED nhấp nháy đỏ

### Quy trình xử lý (theo thứ tự)

**Bước 1 — Kiểm tra nguồn điện**
- Đảm bảo đèn nguồn sáng xanh
- Kiểm tra cáp sạc, thay cáp khác nếu cần
- Charge ít nhất 30 phút nếu pin yếu

**Bước 2 — Kiểm tra kết nối mạng**
- Tắt/bật WiFi router
- Thử kết nối hotspot điện thoại
- Ping test: `ping 8.8.8.8` — nếu timeout là do mạng

**Bước 3 — Khởi động lại thiết bị**
- Giữ nút nguồn 10 giây → đợi màn hình khởi động
- Kiểm tra lại kết nối app

**Bước 4 — Reset về mặc định (nếu vẫn lỗi)**
- Vào Settings → Factory Reset
- ⚠️ Dữ liệu giao dịch local sẽ bị xóa (đã có backup trên server)
- Cấu hình lại merchant ID

**Bước 5 — Escalate**
- Chụp ảnh thiết bị + log lỗi
- Tạo ticket: [support.vng.com.vn](http://support.vng.com.vn)
- Đổi thiết bị dự phòng cho merchant trong lúc chờ xử lý

### Thời gian xử lý cam kết
| Mức độ | Thời gian |
|--------|-----------|
| Bước 1–3 | < 15 phút |
| Bước 4 | < 30 phút |
| Escalate | < 4 giờ làm việc |
"""),
        (4, "Hướng Dẫn Thiết Bị Soundbox SB200", """## Soundbox SB200 — Hướng Dẫn Sử Dụng

### Thông số kỹ thuật
- **Màn hình:** LED 2 dòng, hiển thị số tiền
- **Kết nối:** WiFi 2.4GHz + 4G LTE (SIM backup)
- **Pin:** 2000mAh, nghe thanh toán liên tục 8h
- **Âm lượng:** 85dB, phạm vi nghe 5m
- **Sạc:** USB-C, đầy pin sau 2h

### Các đèn LED

| Màu | Ý nghĩa |
|-----|---------|
| Xanh nhấp nháy | Đang kết nối mạng |
| Xanh liên tục | Kết nối tốt, sẵn sàng |
| Đỏ nhấp nháy | Mất kết nối / lỗi |
| Vàng | Pin yếu (< 20%) |
| Tắt | Nguồn off hoặc hết pin |

### Thao tác cơ bản

**Bật/Tắt:** Giữ nút nguồn 3 giây

**Tăng/giảm âm lượng:** Nút + / - trên thân máy (10 mức)

**Kiểm tra kết nối:** Nhấn nút nguồn 1 lần → LED nhấp nháy xanh 3 lần = OK

**Xem balance:** Giữ nút + trong 2 giây (nếu có màn hình)

### Quy trình Reset
1. Tắt nguồn hoàn toàn
2. Giữ nút + và nút nguồn cùng lúc 10 giây
3. Màn hình hiển thị "RESET" → thả tay
4. Chờ khởi động lại (~60 giây)
5. Cấu hình lại qua app
"""),
    ]

    for section, title, content in docs:
        c.execute(
            "INSERT INTO documents (section, title, content, file_type) VALUES (?, ?, ?, 'md')",
            (section, title, content)
        )

    # Sample quiz questions
    quizzes = [
        (1, "TPV là viết tắt của từ gì?",
         "Total Payment Volume", "Transaction Per Visit",
         "Transfer Payment Voucher", "Total Purchase Value",
         "A", "TPV = Total Payment Volume — Tổng Doanh Số Thanh Toán"),
        (1, "CR% = 20% có nghĩa là gì?",
         "20% khách hàng hủy hợp đồng", "20% lead chuyển thành merchant active",
         "Doanh thu tăng 20%", "20% giao dịch bị lỗi",
         "B", "Conversion Rate 20% nghĩa là cứ 100 lead thì có 20 merchant kích hoạt thành công"),
        (2, "Quy trình onboarding merchant mới tối đa bao nhiêu ngày?",
         "1 ngày", "3 ngày",
         "5 ngày", "7 ngày",
         "B", "Theo SOP chuẩn, onboarding phải hoàn thành trong tối đa 3 ngày làm việc"),
        (3, "Khi Soundbox không kết nối, bước đầu tiên cần kiểm tra là gì?",
         "Reset về mặc định ngay", "Gọi hotline hỗ trợ",
         "Kiểm tra nguồn điện và cáp sạc", "Đổi thiết bị mới",
         "C", "Luôn kiểm tra nguồn điện trước — đây là nguyên nhân phổ biến nhất"),
        (4, "Đèn LED đỏ nhấp nháy trên Soundbox SB200 có nghĩa là gì?",
         "Pin đầy", "Kết nối tốt",
         "Mất kết nối hoặc lỗi", "Đang sạc pin",
         "C", "Đỏ nhấp nháy = mất kết nối hoặc có lỗi cần xử lý"),
    ]

    for q in quizzes:
        c.execute(
            """INSERT INTO quiz_questions
               (section, question, option_a, option_b, option_c, option_d, correct_answer, explanation)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            q
        )

    log.info("Seeded sample data: %d documents, %d quiz questions", len(docs), len(quizzes))


# ---------------------------------------------------------------------------
# Auth helpers
# ---------------------------------------------------------------------------

def _hash_password(password: str) -> str:
    salt = "ops_agent_salt_v1"
    return hashlib.sha256(f"{salt}{password}".encode()).hexdigest()


def _parse_jwt_expires() -> int:
    """Parse JWT_EXPIRES_IN (e.g. '8h', '30m') to seconds."""
    val = JWT_EXPIRES_IN.strip()
    if val.endswith("h"):
        return int(val[:-1]) * 3600
    elif val.endswith("m"):
        return int(val[:-1]) * 60
    elif val.endswith("d"):
        return int(val[:-1]) * 86400
    return 28800  # default 8h


def create_token(user_id: int, username: str, role: str) -> str:
    now = datetime.utcnow()
    payload = {
        "sub": str(user_id),   # JWT spec: sub should be string
        "username": username,
        "role": role,
        "exp": now + timedelta(seconds=_parse_jwt_expires()),
        "iat": now,
    }
    token = jwt.encode(payload, JWT_SECRET, algorithm="HS256")
    # PyJWT >=2.0 returns str; older versions return bytes — normalise
    return token if isinstance(token, str) else token.decode("utf-8")


def verify_token(token: str) -> dict:
    try:
        payload = jwt.decode(
            token,
            JWT_SECRET,
            algorithms=["HS256"],
            options={"verify_aud": False},
        )
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token đã hết hạn, vui lòng đăng nhập lại")
    except jwt.InvalidTokenError as e:
        log.warning("JWT decode failed: %s", e)
        raise HTTPException(status_code=401, detail="Token không hợp lệ")


def _extract_token(request: Request) -> str:
    """Extract Bearer token from Authorization header (robust, no library quirks)."""
    auth = request.headers.get("authorization") or request.headers.get("Authorization") or ""
    if not auth.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Yêu cầu đăng nhập (thiếu Bearer token)")
    return auth[7:].strip()


async def get_current_user(request: Request) -> dict:
    token = _extract_token(request)
    return verify_token(token)


async def require_admin(request: Request) -> dict:
    user = await get_current_user(request)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Chỉ Admin mới có quyền thực hiện thao tác này")
    return user


# ---------------------------------------------------------------------------
# Tool implementations
# ---------------------------------------------------------------------------

def tool_search_knowledge_base(query: str, section: Optional[int] = None, limit: int = 5) -> dict:
    """FTS5 full-text search across documents."""
    try:
        conn = get_connection()
        c = conn.cursor()

        if section:
            rows = c.execute(
                """SELECT d.id, d.section, d.title,
                          snippet(documents_fts, 1, '<b>', '</b>', '...', 32) AS excerpt,
                          rank
                   FROM documents_fts
                   JOIN documents d ON d.id = documents_fts.rowid
                   WHERE documents_fts MATCH ? AND d.section = ?
                   ORDER BY rank
                   LIMIT ?""",
                (query, section, limit)
            ).fetchall()
        else:
            rows = c.execute(
                """SELECT d.id, d.section, d.title,
                          snippet(documents_fts, 1, '<b>', '</b>', '...', 32) AS excerpt,
                          rank
                   FROM documents_fts
                   JOIN documents d ON d.id = documents_fts.rowid
                   WHERE documents_fts MATCH ?
                   ORDER BY rank
                   LIMIT ?""",
                (query, limit)
            ).fetchall()

        conn.close()

        results = [
            {
                "id": r["id"],
                "section": r["section"],
                "section_name": SECTION_NAMES.get(r["section"], ""),
                "title": r["title"],
                "excerpt": r["excerpt"],
            }
            for r in rows
        ]

        return {
            "found": len(results),
            "query": query,
            "results": results,
            "tip": "Dùng get_document_detail(id) để đọc nội dung đầy đủ" if results else None,
        }
    except Exception as e:
        return {"error": str(e), "found": 0, "results": []}


def tool_get_section_content(section: int) -> dict:
    try:
        conn = get_connection()
        c = conn.cursor()
        rows = c.execute(
            "SELECT id, title, file_type, created_at FROM documents WHERE section = ? ORDER BY id",
            (section,)
        ).fetchall()
        conn.close()

        return {
            "section": section,
            "section_name": SECTION_NAMES.get(section, ""),
            "document_count": len(rows),
            "documents": [
                {"id": r["id"], "title": r["title"], "file_type": r["file_type"], "created_at": r["created_at"]}
                for r in rows
            ]
        }
    except Exception as e:
        return {"error": str(e)}


def tool_get_quiz_questions(section: Optional[int] = None, count: int = 10) -> dict:
    try:
        count = min(count, 20)
        conn = get_connection()
        c = conn.cursor()

        if section:
            rows = c.execute(
                """SELECT id, section, question, option_a, option_b, option_c, option_d
                   FROM quiz_questions WHERE section = ?
                   ORDER BY RANDOM() LIMIT ?""",
                (section, count)
            ).fetchall()
        else:
            rows = c.execute(
                """SELECT id, section, question, option_a, option_b, option_c, option_d
                   FROM quiz_questions
                   ORDER BY RANDOM() LIMIT ?""",
                (count,)
            ).fetchall()

        conn.close()

        questions = [
            {
                "id": r["id"],
                "section": r["section"],
                "section_name": SECTION_NAMES.get(r["section"], ""),
                "question": r["question"],
                "options": {
                    "A": r["option_a"],
                    "B": r["option_b"],
                    "C": r["option_c"],
                    "D": r["option_d"],
                }
            }
            for r in rows
        ]

        return {
            "count": len(questions),
            "section": section,
            "questions": questions,
            "instruction": "Trả lời bằng cách chọn A, B, C, hoặc D. Khi xong, tôi sẽ chấm điểm."
        }
    except Exception as e:
        return {"error": str(e)}


def tool_get_document_detail(document_id: int) -> dict:
    try:
        conn = get_connection()
        c = conn.cursor()
        row = c.execute(
            "SELECT id, section, title, content, file_type, created_at, updated_at FROM documents WHERE id = ?",
            (document_id,)
        ).fetchone()
        conn.close()

        if not row:
            return {"error": f"Không tìm thấy tài liệu ID={document_id}"}

        return {
            "id": row["id"],
            "section": row["section"],
            "section_name": SECTION_NAMES.get(row["section"], ""),
            "title": row["title"],
            "content": row["content"],
            "file_type": row["file_type"],
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
        }
    except Exception as e:
        return {"error": str(e)}


def execute_tool(tool_name: str, tool_input: dict) -> str:
    """Dispatch tool calls and return JSON string results."""
    try:
        if tool_name == "search_knowledge_base":
            result = tool_search_knowledge_base(**tool_input)
        elif tool_name == "get_section_content":
            result = tool_get_section_content(**tool_input)
        elif tool_name == "get_quiz_questions":
            result = tool_get_quiz_questions(**tool_input)
        elif tool_name == "get_document_detail":
            result = tool_get_document_detail(**tool_input)
        else:
            result = {"error": f"Unknown tool: {tool_name}"}
    except Exception as e:
        result = {"error": str(e)}

    return json.dumps(result, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Agentic loop (streaming)
# ---------------------------------------------------------------------------

async def run_agent_stream(
    user_message: str,
    conversation_history: list,
) -> AsyncGenerator[str, None]:
    """
    Run agentic loop with tool use (OpenAI-compatible, GreenNode MaaS).
    Yields server-sent event strings.
    """
    if not API_KEY:
        yield "data: " + json.dumps({"type": "error", "text": "API_KEY chưa được cấu hình. Liên hệ Admin."}) + "\n\n"
        return

    # Chọn SDK dựa trên API_BASE_URL:
    # - Trống hoặc key bắt đầu bằng sk-ant- → Anthropic SDK
    # - Có base_url (GreenNode/OpenAI-compatible) → OpenAI SDK
    use_anthropic = not API_BASE_URL or API_KEY.startswith("sk-ant-")

    if use_anthropic:
        async for chunk in _run_anthropic_stream(user_message, conversation_history):
            yield chunk
    else:
        async for chunk in _run_openai_stream(user_message, conversation_history):
            yield chunk


async def _run_anthropic_stream(
    user_message: str,
    conversation_history: list,
) -> AsyncGenerator[str, None]:
    """Agentic loop using Anthropic SDK (for sk-ant- keys)."""
    # Convert TOOLS from OpenAI format to Anthropic format
    anthropic_tools = [
        {
            "name": t["function"]["name"],
            "description": t["function"]["description"],
            "input_schema": t["function"]["parameters"],
        }
        for t in TOOLS
    ]

    client = anthropic_sdk.Anthropic(api_key=API_KEY)
    messages = list(conversation_history)
    messages.append({"role": "user", "content": user_message})

    max_iterations = 10
    iteration = 0

    while iteration < max_iterations:
        iteration += 1

        response = client.messages.create(
            model=MODEL,
            max_tokens=MAX_TOKENS,
            system=SYSTEM_PROMPT,
            tools=anthropic_tools,
            messages=messages,
        )

        text_content = ""
        tool_use_blocks = []
        for block in response.content:
            if block.type == "text":
                text_content += block.text
            elif block.type == "tool_use":
                tool_use_blocks.append(block)

        if text_content:
            words = text_content.split(" ")
            for i, word in enumerate(words):
                chunk = word + (" " if i < len(words) - 1 else "")
                yield "data: " + json.dumps({"type": "text", "text": chunk}) + "\n\n"
                await asyncio.sleep(0.01)

        if response.stop_reason == "end_turn":
            yield "data: " + json.dumps({"type": "done"}) + "\n\n"
            conversation_history.append({"role": "user", "content": user_message})
            conversation_history.append({"role": "assistant", "content": response.content})
            break

        if response.stop_reason == "tool_use" and tool_use_blocks:
            messages.append({"role": "assistant", "content": response.content})
            for block in tool_use_blocks:
                yield "data: " + json.dumps({"type": "tool_call", "tool": block.name, "input": block.input}) + "\n\n"
                await asyncio.sleep(0.05)
            tool_results = []
            for block in tool_use_blocks:
                result_text = await asyncio.get_event_loop().run_in_executor(
                    None, execute_tool, block.name, block.input
                )
                tool_results.append({"type": "tool_result", "tool_use_id": block.id, "content": result_text})
            messages.append({"role": "user", "content": tool_results})
        else:
            yield "data: " + json.dumps({"type": "done"}) + "\n\n"
            break


async def _run_openai_stream(
    user_message: str,
    conversation_history: list,
) -> AsyncGenerator[str, None]:
    """Agentic loop using OpenAI-compatible SDK (for GreenNode/OpenAI keys)."""
    client = OpenAI(api_key=API_KEY, base_url=API_BASE_URL)
    messages = [{"role": "system", "content": SYSTEM_PROMPT}]
    messages += list(conversation_history)
    messages.append({"role": "user", "content": user_message})

    max_iterations = 10
    iteration = 0

    while iteration < max_iterations:
        iteration += 1

        response = client.chat.completions.create(
            model=MODEL,
            max_tokens=MAX_TOKENS,
            tools=TOOLS,
            messages=messages,
        )

        choice = response.choices[0]
        finish_reason = choice.finish_reason
        msg = choice.message

        if msg.content:
            words = msg.content.split(" ")
            for i, word in enumerate(words):
                chunk = word + (" " if i < len(words) - 1 else "")
                yield "data: " + json.dumps({"type": "text", "text": chunk}) + "\n\n"
                await asyncio.sleep(0.01)

        if finish_reason == "stop":
            yield "data: " + json.dumps({"type": "done"}) + "\n\n"
            conversation_history.append({"role": "user", "content": user_message})
            conversation_history.append({"role": "assistant", "content": msg.content or ""})
            break

        if finish_reason == "tool_calls" and msg.tool_calls:
            messages.append({
                "role": "assistant",
                "content": msg.content,
                "tool_calls": [
                    {"id": tc.id, "type": "function",
                     "function": {"name": tc.function.name, "arguments": tc.function.arguments}}
                    for tc in msg.tool_calls
                ],
            })
            for tc in msg.tool_calls:
                try:
                    tool_input = json.loads(tc.function.arguments)
                except Exception:
                    tool_input = {}
                yield "data: " + json.dumps({"type": "tool_call", "tool": tc.function.name, "input": tool_input}) + "\n\n"
                await asyncio.sleep(0.05)
            for tc in msg.tool_calls:
                try:
                    tool_input = json.loads(tc.function.arguments)
                except Exception:
                    tool_input = {}
                result_text = await asyncio.get_event_loop().run_in_executor(
                    None, execute_tool, tc.function.name, tool_input
                )
                messages.append({"role": "tool", "tool_call_id": tc.id, "content": result_text})
        else:
            yield "data: " + json.dumps({"type": "done"}) + "\n\n"
            break


# ---------------------------------------------------------------------------
# Pydantic models
# ---------------------------------------------------------------------------

class LoginRequest(BaseModel):
    username: str
    password: str


class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None


class DocumentCreate(BaseModel):
    section: int
    title: str
    content: str
    file_type: str = "md"


class QuizAttemptRequest(BaseModel):
    answers: dict  # {question_id: "A"|"B"|"C"|"D"}
    section: Optional[int] = None


class UserCreate(BaseModel):
    username: str
    password: str
    role: str = "user"


# ---------------------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------------------

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup/shutdown lifecycle."""
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    init_db()
    log.info("Ops AI Agent started — model=%s port=%d", MODEL, PORT)
    yield
    log.info("Ops AI Agent shutting down")


app = FastAPI(
    title="Bách Khoa Toàn Thư Ops — AI Agent",
    description="Internal Knowledge Base AI for Merchant Operations Team",
    version=VERSION,
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Frontend — serve single-file SPA
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    """Serve the React-less SPA frontend."""
    if FRONTEND_PATH.exists():
        return HTMLResponse(content=FRONTEND_PATH.read_text(encoding="utf-8"))
    return HTMLResponse(content="<h1>Frontend not found</h1><p>Place frontend.html next to agent.py</p>", status_code=404)


# ---------------------------------------------------------------------------
# Health check (required by GreenNode AgentBase)
# ---------------------------------------------------------------------------

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "version": VERSION,
        "model": MODEL,
        "db": "connected" if Path(get_db_path()).exists() else "not_found",
    }


# ---------------------------------------------------------------------------
# Auth endpoints
# ---------------------------------------------------------------------------

@app.post("/api/auth/login")
async def login(req: LoginRequest):
    conn = get_connection()
    user = conn.execute(
        "SELECT id, username, role FROM users WHERE username = ? AND password_hash = ?",
        (req.username, _hash_password(req.password))
    ).fetchone()
    conn.close()

    if not user:
        raise HTTPException(status_code=401, detail="Sai tên đăng nhập hoặc mật khẩu")

    token = create_token(user["id"], user["username"], user["role"])
    return {
        "token": token,
        "username": user["username"],
        "role": user["role"],
        "expires_in": JWT_EXPIRES_IN,
    }


# ---------------------------------------------------------------------------
# Chat endpoint (streaming)
# ---------------------------------------------------------------------------

@app.post("/chat")
async def chat(req: ChatRequest, user: dict = Depends(get_current_user)):
    """Stream AI response with tool-use agentic loop."""
    session_id = req.session_id or secrets.token_hex(8)

    # Load or create session
    conn = get_connection()
    row = conn.execute(
        "SELECT messages FROM chat_sessions WHERE id = ?", (session_id,)
    ).fetchone()

    if row:
        history = json.loads(row["messages"])
    else:
        history = []
        conn.execute(
            "INSERT INTO chat_sessions (id, user_id, messages) VALUES (?, ?, '[]')",
            (session_id, user["sub"])
        )
        conn.commit()
    conn.close()

    async def stream_and_save():
        collected_response = []
        async for chunk in run_agent_stream(req.message, history):
            yield chunk
            # Collect text chunks
            try:
                data = json.loads(chunk.replace("data: ", "").strip())
                if data.get("type") == "text":
                    collected_response.append(data["text"])
            except Exception:
                pass

        # Save updated history to DB
        final_text = "".join(collected_response)
        history.append({"role": "user", "content": req.message})
        history.append({"role": "assistant", "content": final_text})

        # Keep last 20 turns to avoid DB bloat
        trimmed = history[-40:]
        conn2 = get_connection()
        conn2.execute(
            """INSERT INTO chat_sessions (id, user_id, messages, updated_at)
               VALUES (?, ?, ?, CURRENT_TIMESTAMP)
               ON CONFLICT(id) DO UPDATE SET messages=excluded.messages, updated_at=excluded.updated_at""",
            (session_id, user["sub"], json.dumps(trimmed, ensure_ascii=False))
        )
        conn2.commit()
        conn2.close()

    return StreamingResponse(
        stream_and_save(),
        media_type="text/event-stream",
        headers={
            "X-Session-ID": session_id,
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        }
    )


@app.get("/chat/history/{session_id}")
async def get_chat_history(session_id: str, user: dict = Depends(get_current_user)):
    conn = get_connection()
    row = conn.execute(
        "SELECT messages, created_at, updated_at FROM chat_sessions WHERE id = ? AND user_id = ?",
        (session_id, user["sub"])
    ).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Session không tồn tại")

    return {
        "session_id": session_id,
        "messages": json.loads(row["messages"]),
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }


# ---------------------------------------------------------------------------
# Documents API
# ---------------------------------------------------------------------------

@app.get("/api/docs")
async def list_documents(
    section: Optional[int] = None,
    search: Optional[str] = None,
    limit: int = 20,
    user: dict = Depends(get_current_user)
):
    conn = get_connection()

    if search:
        if section:
            rows = conn.execute(
                """SELECT d.id, d.section, d.title, d.file_type, d.created_at,
                          snippet(documents_fts, 1, '<<', '>>', '...', 20) AS excerpt
                   FROM documents_fts
                   JOIN documents d ON d.id = documents_fts.rowid
                   WHERE documents_fts MATCH ? AND d.section = ?
                   ORDER BY rank LIMIT ?""",
                (search, section, limit)
            ).fetchall()
        else:
            rows = conn.execute(
                """SELECT d.id, d.section, d.title, d.file_type, d.created_at,
                          snippet(documents_fts, 1, '<<', '>>', '...', 20) AS excerpt
                   FROM documents_fts
                   JOIN documents d ON d.id = documents_fts.rowid
                   WHERE documents_fts MATCH ?
                   ORDER BY rank LIMIT ?""",
                (search, limit)
            ).fetchall()
    else:
        if section:
            rows = conn.execute(
                "SELECT id, section, title, file_type, created_at, '' AS excerpt FROM documents WHERE section = ? ORDER BY id LIMIT ?",
                (section, limit)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT id, section, title, file_type, created_at, '' AS excerpt FROM documents ORDER BY id LIMIT ?",
                (limit,)
            ).fetchall()

    conn.close()

    return {
        "total": len(rows),
        "documents": [
            {
                "id": r["id"],
                "section": r["section"],
                "section_name": SECTION_NAMES.get(r["section"], ""),
                "title": r["title"],
                "file_type": r["file_type"],
                "created_at": r["created_at"],
                "excerpt": r["excerpt"],
            }
            for r in rows
        ]
    }


@app.get("/api/docs/{doc_id}")
async def get_document(doc_id: int, user: dict = Depends(get_current_user)):
    conn = get_connection()
    row = conn.execute(
        "SELECT * FROM documents WHERE id = ?", (doc_id,)
    ).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Tài liệu không tồn tại")

    return {
        "id": row["id"],
        "section": row["section"],
        "section_name": SECTION_NAMES.get(row["section"], ""),
        "title": row["title"],
        "content": row["content"],
        "file_type": row["file_type"],
        "filename": row["filename"],
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }


@app.post("/api/docs", status_code=201)
async def create_document(doc: DocumentCreate, user: dict = Depends(require_admin)):
    conn = get_connection()
    cursor = conn.execute(
        "INSERT INTO documents (section, title, content, file_type) VALUES (?, ?, ?, ?)",
        (doc.section, doc.title, doc.content, doc.file_type)
    )
    conn.commit()
    doc_id = cursor.lastrowid
    conn.close()
    return {"id": doc_id, "message": "Tài liệu đã được tạo"}


@app.post("/api/docs/upload", status_code=201)
async def upload_document(
    section: int = Form(...),
    title: str = Form(...),
    file: UploadFile = File(...),
    user: dict = Depends(require_admin)
):
    # Validate file size
    content = await file.read()
    size_mb = len(content) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(400, f"File quá lớn ({size_mb:.1f}MB). Tối đa {MAX_FILE_SIZE_MB}MB")

    # Save file
    ext = Path(file.filename).suffix.lower()
    safe_name = f"{secrets.token_hex(8)}{ext}"
    file_path = UPLOAD_DIR / safe_name
    file_path.write_bytes(content)

    # Extract text content based on file type
    text_content = f"[File: {file.filename}]\n"
    if ext in (".txt", ".md"):
        try:
            text_content += content.decode("utf-8")
        except Exception:
            text_content += "(Binary content — không thể trích xuất văn bản)"
    elif ext == ".pdf":
        try:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = []
                for i, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(f"--- Trang {i} ---\n{page_text}")
            text_content += "\n\n".join(pages) if pages else "(PDF không chứa text có thể trích xuất)"
        except Exception as e:
            log.warning("PDF extraction failed: %s", e)
            text_content += f"(Không thể trích xuất PDF: {e})"
    elif ext in (".docx",):
        try:
            import docx, io
            doc_obj = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc_obj.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            text_content += "\n\n".join(paragraphs) if paragraphs else "(DOCX không có nội dung)"
        except Exception as e:
            log.warning("DOCX extraction failed: %s", e)
            text_content += f"(Không thể trích xuất DOCX: {e})"
    elif ext in (".xlsx", ".xls"):
        try:
            import openpyxl, io
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sheets = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(c) if c is not None else "" for c in row]
                    if any(v.strip() for v in row_vals):
                        rows.append(" | ".join(row_vals))
                if rows:
                    sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
            text_content += "\n\n".join(sheets) if sheets else "(XLSX không có dữ liệu)"
        except Exception as e:
            log.warning("XLSX extraction failed: %s", e)
            text_content += f"(Không thể trích xuất XLSX: {e})"
    else:
        text_content += f"(File {ext.upper()} — {size_mb:.1f}MB. Xem file gốc để biết chi tiết.)"

    conn = get_connection()
    cursor = conn.execute(
        "INSERT INTO documents (section, title, content, file_type, filename) VALUES (?, ?, ?, ?, ?)",
        (section, title, text_content, ext.lstrip(".") or "bin", safe_name)
    )
    conn.commit()
    conn.close()

    return {"id": cursor.lastrowid, "filename": safe_name, "size_mb": round(size_mb, 2)}


@app.delete("/api/docs/{doc_id}")
async def delete_document(doc_id: int, user: dict = Depends(require_admin)):
    conn = get_connection()
    row = conn.execute("SELECT filename FROM documents WHERE id = ?", (doc_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "Tài liệu không tồn tại")

    conn.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    conn.commit()
    conn.close()

    # Remove physical file if exists
    if row["filename"]:
        try:
            (UPLOAD_DIR / row["filename"]).unlink(missing_ok=True)
        except Exception:
            pass

    return {"message": "Đã xóa tài liệu"}


# ---------------------------------------------------------------------------
# Quiz API
# ---------------------------------------------------------------------------

@app.get("/api/quiz")
async def get_quiz(
    section: Optional[int] = None,
    count: int = 10,
    user: dict = Depends(get_current_user)
):
    result = tool_get_quiz_questions(section=section, count=count)
    return result


@app.post("/api/quiz/attempt")
async def submit_quiz_attempt(req: QuizAttemptRequest, user: dict = Depends(get_current_user)):
    if not req.answers:
        raise HTTPException(400, "Không có câu trả lời nào")

    conn = get_connection()
    question_ids = list(req.answers.keys())
    placeholders = ",".join(["?"] * len(question_ids))
    rows = conn.execute(
        f"SELECT id, question, correct_answer, explanation FROM quiz_questions WHERE id IN ({placeholders})",
        question_ids
    ).fetchall()
    conn.close()

    results = []
    correct = 0
    for row in rows:
        qid = str(row["id"])
        user_ans = req.answers.get(qid, "").upper()
        is_correct = user_ans == row["correct_answer"]
        if is_correct:
            correct += 1
        results.append({
            "id": row["id"],
            "question": row["question"],
            "your_answer": user_ans,
            "correct_answer": row["correct_answer"],
            "is_correct": is_correct,
            "explanation": row["explanation"],
        })

    total = len(rows)
    score = correct
    pct = round(score / total * 100) if total else 0

    if pct >= 90:
        grade = "Xuất sắc"
    elif pct >= 70:
        grade = "Khá"
    elif pct >= 50:
        grade = "Trung bình"
    else:
        grade = "Cần ôn thêm"

    # Save attempt
    conn2 = get_connection()
    conn2.execute(
        "INSERT INTO quiz_attempts (user_id, section, score, total) VALUES (?, ?, ?, ?)",
        (user.get("id"), req.section, score, total)
    )
    conn2.commit()
    conn2.close()

    return {"score": score, "total": total, "percentage": pct, "grade": grade, "results": results}


@app.get("/api/quiz/history")
async def get_quiz_history(user: dict = Depends(get_current_user)):
    conn = get_connection()
    rows = conn.execute(
        "SELECT id, section, score, total, taken_at FROM quiz_attempts WHERE user_id = ? ORDER BY taken_at DESC LIMIT 20",
        (user.get("id"),)
    ).fetchall()
    conn.close()
    return {"attempts": [dict(r) for r in rows]}


# ---------------------------------------------------------------------------
# Admin API
# ---------------------------------------------------------------------------

@app.get("/api/stats")
async def get_stats(user: dict = Depends(require_admin)):
    conn = get_connection()
    docs_count = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
    quiz_count = conn.execute("SELECT COUNT(*) FROM quiz_questions").fetchone()[0]
    attempt_count = conn.execute("SELECT COUNT(*) FROM quiz_attempts").fetchone()[0]
    user_count = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    conn.close()
    return {
        "documents": docs_count,
        "quiz_questions": quiz_count,
        "quiz_attempts": attempt_count,
        "users": user_count,
    }


@app.get("/api/admin/users")
async def list_users(user: dict = Depends(require_admin)):
    conn = get_connection()
    rows = conn.execute(
        "SELECT id, username, role, created_at FROM users ORDER BY created_at DESC"
    ).fetchall()
    conn.close()
    return {"users": [dict(r) for r in rows]}


@app.post("/api/admin/users", status_code=201)
async def create_user(req: UserCreate, user: dict = Depends(require_admin)):
    if req.role not in ("admin", "user"):
        raise HTTPException(400, "Role không hợp lệ")
    pw_hash = hashlib.sha256(req.password.encode()).hexdigest()
    try:
        conn = get_connection()
        cursor = conn.execute(
            "INSERT INTO users (username, password_hash, role) VALUES (?, ?, ?)",
            (req.username.strip(), pw_hash, req.role)
        )
        conn.commit()
        conn.close()
        return {"id": cursor.lastrowid, "username": req.username, "role": req.role}
    except Exception:
        raise HTTPException(409, "Username đã tồn tại")


@app.delete("/api/admin/users/{user_id}")
async def delete_user(user_id: int, user: dict = Depends(require_admin)):
    if user.get("id") == user_id:
        raise HTTPException(400, "Không thể xóa tài khoản của chính mình")
    conn = get_connection()
    row = conn.execute("SELECT id FROM users WHERE id = ?", (user_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "Người dùng không tồn tại")
    conn.execute("DELETE FROM users WHERE id = ?", (user_id,))
    conn.commit()
    conn.close()
    return {"message": "Đã xóa người dùng"}


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("agent:app", host="0.0.0.0", port=PORT, reload=False)
